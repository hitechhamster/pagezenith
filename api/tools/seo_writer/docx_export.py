"""Markdown 文章 → Word（.docx）字节流。

移植自线下工作流：标题层级、无序列表、表格、加粗、Markdown 超链接（真·可点蓝色下划线），
以及把正文里的 `[IMAGE: ...]` 占位符替换成生成好的图片。

不落盘：直接返回 bytes，由 router 以 base64 随 SSE done 事件发给浏览器下载。
"""

from __future__ import annotations

import io
import re
from typing import Optional

import docx
import docx.enum.text
import docx.opc.constants
import docx.shared
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

IMAGE_PLACEHOLDER = re.compile(r"^\[IMAGE:\s*[^\]]+?\s*\]$")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """python-docx 没有原生超链接 API，只能手搓 w:hyperlink 元素。"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _add_formatted_runs(paragraph, text: str) -> None:
    """把一行 Markdown 拆成 [链接 / 加粗 / 普通] 三种 run。"""
    for part in re.split(r"(\[.*?\]\(.*?\)|\*\*.*?\*\*)", text):
        link = re.match(r"\[(.*?)\]\((.*?)\)", part)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2))
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


# ── Markdown 表格识别 ───────────────────────────────────────────────
# 模型两种写法都会吐：标准（|a|b|）和松散（a | b）。只认标准的话，松散表格会整块
# 变成正文里的一堆管道符。判据改成「这一行有竖线，且**下一行是分隔行**」——
# 分隔行才是 Markdown 表格真正的标志。（md.js 里的前端渲染器同款逻辑。）
_SEP_ROW = re.compile(r"^[|\s:\-]+$")


def _is_sep_row(line: str) -> bool:
    t = (line or "").strip()
    return bool(t) and "|" in t and bool(_SEP_ROW.match(t)) and "--" in t


def _is_table_head(line: str, next_line: str) -> bool:
    return "|" in (line or "") and _is_sep_row(next_line)


def _split_row(row_text: str) -> list:
    """按竖线切一行，首尾的空格和竖线都容忍。"""
    return [c.strip() for c in row_text.strip().strip("|").split("|")]


def build_docx(markdown_text: str, image_map: Optional[dict[str, bytes]] = None) -> bytes:
    """返回 .docx 的字节内容。image_map: {"[IMAGE: xxx]": png_bytes}"""
    doc = docx.Document()
    image_map = image_map or {}

    lines = (markdown_text or "").strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 图片占位符 → 真图
        if IMAGE_PLACEHOLDER.match(line):
            png = image_map.get(line)
            if png:
                try:
                    doc.add_picture(io.BytesIO(png), width=docx.shared.Inches(6.0))
                    doc.paragraphs[-1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass  # 单张图插不进去不该让整篇导出失败
            i += 1
            continue

        # Markdown 表格
        #
        # ⚠️ 2026-09-04 修：原本要求每行**首尾都有 `|`** 才认。但模型经常吐松散写法
        #    （`App | Price | Best for` 配 `--- | --- | ---`，两头不带竖线），那种情况下
        #    整张表会原样落进 Word 变成一堆管道符。前端渲染器 md.js 已经踩过同一个坑，
        #    这里是同款修法：**靠「下一行是分隔行」来识别表头**，不再要求首尾竖线。
        if _is_table_head(line, lines[i + 1] if i + 1 < len(lines) else ""):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1
            data_rows = [l for l in table_lines if not _SEP_ROW.match(l)]
            if not data_rows:
                continue
            header_cols = _split_row(data_rows[0])
            if not header_cols:
                continue
            table = doc.add_table(rows=1, cols=len(header_cols))
            table.style = "Table Grid"
            for j, h in enumerate(header_cols):
                p = table.rows[0].cells[j].paragraphs[0]
                _add_formatted_runs(p, h)
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for row_text in data_rows[1:]:
                cells = _split_row(row_text)
                row = table.add_row().cells
                for j, ct in enumerate(cells):
                    if j < len(row):
                        _add_formatted_runs(row[j].paragraphs[0], ct)
            continue

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            _add_formatted_runs(doc.add_heading(level=1), line[2:].strip())
        elif line.startswith("## "):
            _add_formatted_runs(doc.add_heading(level=2), line[3:].strip())
        elif line.startswith("### "):
            _add_formatted_runs(doc.add_heading(level=3), line[4:].strip())
        elif line.startswith("#### "):
            _add_formatted_runs(doc.add_heading(level=4), line[5:].strip())
        elif line.startswith("* ") or line.startswith("- "):
            _add_formatted_runs(doc.add_paragraph(style="List Bullet"), line[2:].strip())
        else:
            _add_formatted_runs(doc.add_paragraph(), line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name or "").strip()[:100] or "article"
